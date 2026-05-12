"""
Regenera Anexo_Figuras_Capturas.docx incorporando:
  - Las figuras web-admin (F20..F47, ya existentes).
  - La nueva sección móvil con las capturas REALES del emulador Android
    (RFM01..RFM22) que documentan el flujo end-to-end de entrega:
        login → mis rutas → pedidos reales MICOTRANS → cambio de estado
        → foto evidencia → DNI receptor → firma digital → confirmación.
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

HERE = Path(__file__).resolve().parent
FIG  = HERE / "figuras_capturas"
OUT  = HERE / "docx" / "Anexo_Figuras_Capturas.docx"

# Web-admin (alta horizontal): ancho razonable para landscape
WEB_FIGS = [
    ("F20_login.png",              "Pantalla de Login del Panel Administrativo EcoRoute."),
    ("F21_home.png",               "Pantalla Home / Bienvenida al sistema."),
    ("F22_dashboard.png",          "Panel de Control Logístico con KPIs operativos."),
    ("F23_pedidos.png",            "Gestión de Pedidos mostrando los datos reales de MICOTRANS (guías GR-1001 a GR-1024)."),
    ("F24_rutas.png",              "Planificación y Seguimiento de Rutas con mapa y polilínea."),
    ("F25_conductores.png",        "Gestión de los 5 conductores asignados a MICOTRANS."),
    ("F26_vehiculos.png",          "Gestión de la flota AFT-101 a AFT-105 de MICOTRANS."),
    ("F39_dashboard_top.png",      "Dashboard en ejecución mostrando los pedidos del día."),
    ("F47_kpi_dashboard_post.png", "Dashboard de KPIs de Gestión Administrativa (Tesis) - vista Post-Test."),
    ("F47b_kpi_dashboard_pre.png", "Dashboard de KPIs de Gestión Administrativa (Tesis) - vista Pre-Test."),
]

# Móvil (capturas reales Pixel 9 Pro XL 1344x2992): ancho vertical reducido
MOBILE_FLOW = [
    ("RFM01_login_inicial.png",
     "Pantalla inicial tras iniciar la aplicación EcoRoute Driver en el emulador Android (Pixel 9 Pro XL). "
     "La sesión del conductor (usuario «conductor») persiste y carga directamente «Mis Rutas de Hoy»."),
    ("RFM02_rutas_cargadas.png",
     "«Mis Rutas de Hoy» con datos reales de MICOTRANS S.A.C.: Ruta #4 del 2026-03-05 con seis entregas pendientes "
     "(GR-1019 a GR-1024). El mapa muestra los puntos de entrega georreferenciados en Lima."),
    ("RFM03_detalle_GR1020_PENDING.png",
     "Detalle del pedido GR-1020 (cliente Metal Mecánica El Pino — Victoria) en estado PENDING. "
     "El conductor puede cambiar el estado, adjuntar evidencia fotográfica y guardar."),
    ("RFM04_dropdown_abierto.png",
     "Selector de estados desplegado mostrando las cuatro transiciones permitidas: "
     "PENDING, IN_TRANSIT, DELIVERED y FAILED."),
    ("RFM05_estado_in_transit.png",
     "Estado cambiado a IN_TRANSIT (resaltado en verde). La UI confirma visualmente la selección."),
    ("RFM06_camara_abierta.png",
     "Cámara del dispositivo abierta como parte del flujo de evidencia fotográfica. "
     "(La textura mostrada corresponde a la escena virtual del emulador Android)."),
    ("RFM07_foto_tomada.png",
     "Pantalla de revisión post-disparo. El operador puede aceptar (✓), repetir (↩) o descartar (✗) la captura."),
    ("RFM08_evidencia_guardada.png",
     "Foto aceptada y embebida en el formulario de detalle como evidencia fotográfica. El botón GUARDAR ESTADO queda habilitado."),
    ("RFM09_guardado_in_transit.png",
     "Tras pulsar GUARDAR ESTADO, la app retorna a «Mis Rutas de Hoy» con el toast verde "
     "«Estado actualizado exitosamente». El pedido GR-1020 ya no aparece como PENDING."),
    ("RFM10_lista_post_in_transit.png",
     "Lista refrescada: el pedido GR-1020 aparece reordenado al final como IN_TRANSIT, "
     "evidenciando que el cambio se persistió en el backend."),
    ("RFM11_detalle_in_transit.png",
     "Re-apertura del pedido GR-1020 mostrando el estado persistido IN_TRANSIT."),
    ("RFM12_dropdown_para_delivered.png",
     "Selector desplegado nuevamente, con IN_TRANSIT resaltado como estado actual y DELIVERED disponible para confirmar la entrega."),
    ("RFM13_estado_delivered.png",
     "Selección de DELIVERED. La UI expande dinámicamente los campos obligatorios para una entrega completa: "
     "datos del receptor (nombre + DNI) y firma digital."),
    ("RFM14_evidencia_delivered.png",
     "Foto de evidencia tomada nuevamente para el estado DELIVERED. El nombre del receptor se pre-llena con el cliente del pedido."),
    ("RFM15_dni_ingresado.png",
     "DNI del receptor ingresado (8/8 caracteres validados). Teclado numérico nativo de Android."),
    ("RFM16_firma_canvas.png",
     "Sección «Firma del Receptor» con canvas vacío, botones «Limpiar» y «Confirmar Firma», y «GUARDAR ESTADO» al pie."),
    ("RFM17_firma_dibujada.png",
     "Canvas con firma dibujada (trazos cruzados). El SDK Flutter captura el path como SVG y lo serializa a PNG."),
    ("RFM18_firma_confirmada.png",
     "Confirmación de firma: aparece «✓ Firma guardada» en verde junto al canvas. La firma queda lista para subir como parte de la prueba de entrega."),
    ("RFM19_error_upload_evidencia.png",
     "Evidencia de manejo de errores: cuando el bucket S3 del sistema de almacenamiento no responde, "
     "la app muestra un toast rojo «Fallo de red al subir evidencia». El usuario no pierde los datos ingresados."),
    ("RFM22_guardado_delivered_final.png",
     "Lista final de «Mis Rutas de Hoy» tras la entrega exitosa: GR-1020 figura ahora como DELIVERED. "
     "Foto y firma se almacenan en S3 (bucket ecoroute-proofs/photos|signatures) y se publica un evento en la cola SQS ecoroute-notifications."),
    ("RFM_extra_entrega_finalizada.png",
     "Pantalla de control: al abrir un pedido ya entregado (GR-1021), el aplicativo bloquea la modificación con el aviso "
     "«ENTREGA FINALIZADA - Este pedido ya no puede ser modificado», preservando la integridad del registro de entrega."),
]


def set_run_size(p, size_pt=10, italic=False, bold=False, color=None):
    for r in p.runs:
        r.font.size = Pt(size_pt)
        r.italic = italic
        r.bold = bold
        if color:
            r.font.color.rgb = color


def add_figure(doc, img_path: Path, fig_num: str, caption: str, width_in: float):
    """Insertar una figura con número de figura como heading y caption en italic."""
    h = doc.add_heading(f"Figura {fig_num}", level=3)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    p_cap = doc.add_paragraph(f"Figura {fig_num}: {caption}")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_size(p_cap, size_pt=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))


def main():
    doc = Document()
    # ----- Title -----
    title = doc.add_heading("Anexo Visual — Capturas del Sistema EcoRoute", level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "Este anexo recopila las capturas reales del sistema EcoRoute en operación tomadas durante "
        "las pruebas con datos auténticos de MICOTRANS S.A.C. (24 guías de remisión históricas, "
        "5 conductores, 5 vehículos de la flota AFT-101..AFT-105). Las capturas web fueron "
        "tomadas a 2880×1800; las capturas móviles, en el emulador Android Pixel 9 Pro XL "
        "(1344×2992) ejecutando el binario Flutter compilado contra el backend Spring Boot, "
        "PostgreSQL, LocalStack (S3 + SQS) y Keycloak desplegados con docker-compose."
    )

    # ----- Sección 1: Panel Web Administrativo -----
    doc.add_heading("1. Panel Web Administrativo (web-admin)", level=2)
    fig_counter = 20
    for fname, caption in WEB_FIGS:
        path = FIG / fname
        if not path.exists():
            print(f"  WARN: falta {fname}")
            continue
        add_figure(doc, path, str(fig_counter), caption, width_in=6.2)
        fig_counter += 1

    # ----- Sección 2: Aplicativo Móvil (Conductor) -----
    doc.add_heading("2. Aplicativo Móvil del Conductor — Flujo End-to-End de Entrega", level=2)
    intro_m = doc.add_paragraph(
        "Esta sección documenta paso a paso el flujo completo de entrega ejecutado sobre el emulador Android "
        "(usuario conductor / conductor123) contra los datos reales de la base de MICOTRANS. Se muestra la "
        "transición de un pedido pendiente al estado IN_TRANSIT, y posteriormente a DELIVERED con foto, DNI "
        "del receptor y firma digital. Incluye la prueba negativa de la regla de negocio que bloquea la "
        "modificación de entregas ya finalizadas."
    )

    fig_counter = 50
    for fname, caption in MOBILE_FLOW:
        path = FIG / fname
        if not path.exists():
            print(f"  WARN: falta {fname}")
            continue
        # Vertical 1344x2992 → ancho razonable 2.6 in para que quepan dos por página
        add_figure(doc, path, f"M-{fig_counter}", caption, width_in=2.8)
        fig_counter += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")
    print(f"  Tamano: {OUT.stat().st_size//1024} KB")


if __name__ == "__main__":
    main()
