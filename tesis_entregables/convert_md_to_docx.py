"""
Convierte todos los archivos .md de la carpeta tesis_entregables/ a .docx
usando python-docx. Soporta:
- Headings (#, ##, ###)
- Tablas markdown
- Listas con viñetas y numeradas
- Texto en negrita (**) e itálica (*)
- Bloques de código ```
- Citas (>)

USO: python convert_md_to_docx.py
SALIDA: tesis_entregables/docx/*.docx
"""
import re
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THIS_DIR = Path(__file__).parent
OUT_DIR = THIS_DIR / "docx"
OUT_DIR.mkdir(exist_ok=True)


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_runs_from_inline(paragraph, text):
    """Procesa texto inline: **bold**, *italic*, `code`."""
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def parse_md_table(lines, start):
    """Parsea una tabla markdown a partir de la línea start. Devuelve (rows, end_idx)."""
    rows = []
    i = start
    while i < len(lines):
        ln = lines[i].strip()
        if not ln or not ln.startswith('|'):
            break
        # Separator row?
        if re.match(r'^\|[\s\-:|]+\|$', ln):
            i += 1
            continue
        cells = [c.strip() for c in ln.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def convert_md(md_path: Path, docx_path: Path):
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    doc = Document()
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = text.split('\n')
    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Code blocks
        if stripped.startswith('```'):
            if in_code:
                # Cerrar bloque
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run('\n'.join(code_buf))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=4)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', stripped):
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        # Imágenes markdown: ![alt](path)
        img_match = re.match(r'^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped)
        if img_match:
            alt, img_path = img_match.group(1), img_match.group(2)
            # Resolver path relativo al markdown
            full_path = (THIS_DIR / img_path) if not Path(img_path).is_absolute() else Path(img_path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if full_path.exists():
                try:
                    p.add_run().add_picture(str(full_path), width=Inches(6.0))
                except Exception as e:
                    p.add_run(f"[Imagen no embebida: {img_path} — {e}]")
            else:
                p.add_run(f"[Imagen no encontrada: {img_path}]")
            i += 1
            continue

        # Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            rows, end = parse_md_table(lines, i)
            if len(rows) >= 1:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = 'Light Grid Accent 1'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        if c_idx < len(rows[0]):
                            tcell = tbl.cell(r_idx, c_idx)
                            tcell.text = ''
                            p = tcell.paragraphs[0]
                            if r_idx == 0:
                                set_cell_background(tcell, '2563EB')
                                run = p.add_run(cell)
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            else:
                                add_runs_from_inline(p, cell)
                doc.add_paragraph('')  # espacio después tabla
            i = end
            continue

        # Blockquotes
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
            i += 1
            continue

        # Bullet list
        if stripped.lstrip().startswith('- ') or stripped.lstrip().startswith('* '):
            text_item = stripped.lstrip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_runs_from_inline(p, text_item)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\s*\d+\.\s+', stripped):
            content = re.sub(r'^\s*\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            add_runs_from_inline(p, content)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_runs_from_inline(p, stripped)
        i += 1

    doc.save(str(docx_path))


def main():
    md_files = sorted(THIS_DIR.glob('*.md'))
    print(f"Encontrados {len(md_files)} archivos .md")
    print(f"Directorio salida: {OUT_DIR}")
    print("-" * 60)

    for md in md_files:
        docx = OUT_DIR / f"{md.stem}.docx"
        try:
            convert_md(md, docx)
            size_kb = docx.stat().st_size / 1024
            print(f"  OK  {md.name:60} -> {docx.name} ({size_kb:.1f} KB)")
        except Exception as e:
            print(f"  ERR {md.name}: {e}")

    print("-" * 60)
    print(f"Conversion completada. {len(md_files)} archivos en {OUT_DIR}")


if __name__ == '__main__':
    main()
