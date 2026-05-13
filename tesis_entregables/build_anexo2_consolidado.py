"""
Construye el Anexo 2 CONSOLIDADO del PDF de la tesis:
- Cuestionario UTAUT (instrumento)
- 18 respuestas (tabla)
- Las 6 fichas REALES generadas por el sistema EcoRoute como evidencia visual

Cada PDF de ficha se convierte a PNG y se embebe en el docx.
"""
import sys
import csv
from pathlib import Path
import pypdfium2 as pdfium
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding='utf-8')

ROOT = Path(__file__).parent
FICHAS = ROOT / "fichas_kpi_pdf"
RESPUESTAS_CSV = ROOT / "Anexo_2_Cuestionario_Respuestas.csv"
OUT = ROOT / "docx" / "Anexo_2_Consolidado.docx"
PNG_DIR = ROOT / "fichas_kpi_pdf" / "_png"
PNG_DIR.mkdir(exist_ok=True)


def pdf_to_pngs(pdf_path: Path, scale: float = 2.0):
    """Convierte un PDF a una lista de paths PNG (uno por página)."""
    pdf = pdfium.PdfDocument(str(pdf_path))
    out_paths = []
    for i in range(len(pdf)):
        page = pdf[i]
        pil = page.render(scale=scale).to_pil()
        out_path = PNG_DIR / f"{pdf_path.stem}_p{i+1}.png"
        pil.save(str(out_path), "PNG", optimize=True)
        out_paths.append(out_path)
    return out_paths


def add_centered_image(doc, img_path: Path, width_inches: float = 6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def main():
    doc = Document()

    # Margins razonables
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============== TÍTULO ==============
    title = doc.add_heading("Anexo 2 — Instrumentos de Recolección de Datos", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        "Este anexo contiene el instrumento principal de recolección de datos de la presente "
        "investigación: el cuestionario UTAUT aplicado al personal de MICOTRANS, las respuestas "
        "registradas, y las seis fichas de registro (Pre-Test y Post-Test) de los indicadores "
        "IID, CHR y TDE, generadas automáticamente por el sistema EcoRoute desde su endpoint "
    )
    code = intro.add_run("/reports/kpi/{indicador}/pdf")
    code.font.name = "Consolas"
    code.font.size = Pt(10)
    code.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    intro.add_run(".")

    # ============== PARTE 1: CUESTIONARIO UTAUT ==============
    doc.add_heading("Parte 1 — Cuestionario UTAUT (Instrumento)", level=2)

    p = doc.add_paragraph()
    p.add_run("Tipo de instrumento: ").bold = True
    p.add_run("Cuestionario en escala de Likert de 5 puntos. ")
    p.add_run("Marco teórico: ").bold = True
    p.add_run("UTAUT — Teoría Unificada de Aceptación y Uso de Tecnología. ")
    p.add_run("Muestra: ").bold = True
    p.add_run("18 personas (5 conductores, 6 administrativos, 4 operaciones, 3 gerencia).")

    # Tabla de escala
    doc.add_heading("Escala de medición", level=3)
    t = doc.add_table(rows=6, cols=2)
    t.style = 'Light Grid Accent 1'
    t.cell(0, 0).text = "Valor"
    t.cell(0, 1).text = "Etiqueta"
    items = [("1", "Totalmente en desacuerdo"), ("2", "En desacuerdo"),
             ("3", "Ni de acuerdo ni en desacuerdo"), ("4", "De acuerdo"), ("5", "Totalmente de acuerdo")]
    for i, (v, lab) in enumerate(items, 1):
        t.cell(i, 0).text = v
        t.cell(i, 1).text = lab

    # Preguntas
    doc.add_heading("Preguntas del cuestionario", level=3)
    preguntas = [
        ("Expectativa de Rendimiento", [
            "P1. El aplicativo móvil EcoRoute me permite registrar pedidos con mayor rapidez y exactitud que el método anterior.",
            "P2. El aplicativo móvil EcoRoute mejora mi eficiencia laboral diaria.",
            "P3. Usar EcoRoute ha reducido los errores en mis registros administrativos.",
        ]),
        ("Expectativa de Esfuerzo", [
            "P4. Aprender a usar EcoRoute fue fácil para mí.",
            "P5. La interfaz del aplicativo móvil EcoRoute es clara e intuitiva.",
            "P6. Realizar tareas en el aplicativo (foto, firma) no me toma esfuerzo adicional.",
        ]),
        ("Influencia Social", [
            "P7. Mis compañeros opinan que debería seguir usando el aplicativo EcoRoute.",
            "P8. La gerencia de Grupo MICOTRANS apoya el uso del aplicativo móvil.",
        ]),
        ("Condiciones Facilitadoras", [
            "P9. Cuento con el equipo (celular, datos móviles) necesario para usar EcoRoute.",
            "P10. Recibí la capacitación necesaria para usar el aplicativo correctamente.",
        ]),
        ("Intención y Satisfacción", [
            "P11. Recomendaría el uso del aplicativo móvil EcoRoute a otras empresas del sector transporte.",
            "P12. Considero que el aplicativo móvil agiliza la entrega de evidencias digitales.",
            "P13. En general, estoy satisfecho(a) con el aplicativo móvil EcoRoute.",
        ]),
    ]
    for dim, qs in preguntas:
        h = doc.add_paragraph()
        r = h.add_run(dim)
        r.bold = True
        r.font.size = Pt(11)
        for q in qs:
            doc.add_paragraph(q, style="List Bullet")

    # ============== PARTE 2: RESPUESTAS ==============
    doc.add_page_break()
    doc.add_heading("Parte 2 — Respuestas del Cuestionario (n=18)", level=2)
    intro2 = doc.add_paragraph()
    intro2.add_run(
        "La siguiente tabla contiene las respuestas de los 18 participantes, "
        "incluyendo sus KPIs personales observados durante el post-test."
    )

    if RESPUESTAS_CSV.exists():
        with open(RESPUESTAS_CSV, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = list(reader)
        # Imprimir tabla resumida (id, rol, P01..P13, KPIs)
        cols_show = [0, 1, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20]
        table = doc.add_table(rows=len(rows), cols=len(cols_show))
        table.style = 'Light Grid Accent 1'
        for r_idx, row in enumerate(rows):
            for c_idx, col_idx in enumerate(cols_show):
                if col_idx < len(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    if r_idx == 0:
                        run = p.add_run(row[col_idx])
                        run.bold = True
                        run.font.size = Pt(8)
                    else:
                        run = p.add_run(row[col_idx])
                        run.font.size = Pt(8)

        # Resumen estadístico
        doc.add_paragraph()
        sum_p = doc.add_paragraph()
        sum_p.add_run("Resumen estadístico (calculado con script analisis_estadistico.py):").bold = True
        for line in [
            "  • Alfa de Cronbach (13 ítems) = 0.877 — Buena consistencia interna",
            "  • Pearson r (EE ↔ IID) = 0.927, p < 0.001",
            "  • Pearson r (PE ↔ CHR) = 0.848, p < 0.001",
            "  • Pearson r (Intención ↔ TDE) = 0.736, p < 0.001",
            "  • Pearson r (Satisfacción global ↔ KPI promedio) = 0.984, p < 0.001",
        ]:
            doc.add_paragraph(line)

    # ============== PARTE 3: FICHAS DE REGISTRO ==============
    doc.add_page_break()
    doc.add_heading("Parte 3 — Fichas de Registro generadas por el sistema EcoRoute", level=2)
    p3 = doc.add_paragraph()
    p3.add_run(
        "Las siguientes seis fichas fueron generadas automáticamente por el sistema EcoRoute "
        "desde su endpoint REST "
    )
    cd = p3.add_run("/reports/kpi/{indicador}/pdf")
    cd.font.name = "Consolas"
    cd.font.size = Pt(10)
    p3.add_run(
        ". Cada ficha contiene los datos día por día y el total acumulado, en formato consistente "
        "con el Anexo 2 original de la tesis. Las fichas Pre-Test corresponden a los 150 registros "
        "manuales reales de MICOTRANS S.A.C. (archivo CSV original). Las fichas Post-Test corresponden "
        "a la operación del sistema digital."
    )

    # Mapa de fichas en orden lógico
    fichas_order = [
        ("ficha_iid_pre-test.pdf", "Ficha 1: IID — Pre-Test (registro manual MICOTRANS, IID = 60.0%)"),
        ("ficha_iid_post-test.pdf", "Ficha 2: IID — Post-Test (con sistema EcoRoute, IID = 97.8%)"),
        ("ficha_chr_pre-test.pdf", "Ficha 3: CHR — Pre-Test (registro manual, CHR = 67.3%)"),
        ("ficha_chr_post-test.pdf", "Ficha 4: CHR — Post-Test (con sistema, CHR = 93.9%)"),
        ("ficha_tde_pre-test.pdf", "Ficha 5: TDE — Pre-Test (registro manual, TDE = 51.3%)"),
        ("ficha_tde_post-test.pdf", "Ficha 6: TDE — Post-Test (con sistema, TDE = 93.9%)"),
    ]

    for pdf_name, caption in fichas_order:
        pdf_path = FICHAS / pdf_name
        if not pdf_path.exists():
            print(f"WARN: no existe {pdf_path}")
            continue
        doc.add_page_break()
        h = doc.add_heading(caption.split(":")[0], level=3)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)
        # Convertir PDF a PNG y embeber cada página
        pngs = pdf_to_pngs(pdf_path, scale=2.0)
        for png in pngs:
            add_centered_image(doc, png, width_inches=6.5)

    doc.save(str(OUT))
    print(f"\nOK: {OUT}")
    print(f"Tamaño: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
