"""
Genera un docx dedicado con las 12 figuras del Capítulo III Resultados.

Útil para entregar los gráficos por separado (al asesor, al jurado, o para
copy/paste en la tesis). Cada figura ocupa ~su propia página con:
  - Título numerado (Figura N).
  - La imagen centrada.
  - Pie con interpretación breve.
  - "Fuente: Elaboración propia".

Salida: tesis_entregables/docx/Graficos_Cap3_Resultados.docx
"""
from __future__ import annotations
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THIS_DIR = Path(__file__).parent
FIG_DIR  = THIS_DIR / "figuras_capturas"
OUT      = THIS_DIR / "docx" / "Graficos_Cap3_Resultados.docx"


# Plan de figuras: (num, archivo, título, descripción / pie)
FIGURAS = [
    # ───── Sección 1: Análisis Descriptivo (barras 3D Pre vs Post) ─────
    (14, "FIG_14_BARRAS_IID.png",
     "Figura 14. Comparativo Pre-Test vs Post-Test — IID",
     "Indicador IID — Integridad de Datos Registrados. La media pasó de "
     "58.72% (pre-test, gestión manual) a 97.52% (post-test, con aplicativo), "
     "evidenciando un incremento absoluto de +38.80 puntos porcentuales."),

    (15, "FIG_15_BARRAS_CHR.png",
     "Figura 15. Comparativo Pre-Test vs Post-Test — CHR",
     "Indicador CHR — Cumplimiento de Hoja de Ruta. La media pasó de "
     "65.70% (pre-test) a 93.55% (post-test), incremento absoluto de "
     "+27.85 puntos porcentuales."),

    (16, "FIG_16_BARRAS_TDE.png",
     "Figura 16. Comparativo Pre-Test vs Post-Test — TDE",
     "Indicador TDE — Tasa de Disponibilidad de Evidencias Digitales. "
     "La media pasó de 51.74% a 93.35%, incremento de +41.61 puntos "
     "porcentuales (el mayor de los tres indicadores)."),

    # ───── Sección 2: Histogramas con curva normal (Shapiro-Wilk) ─────
    (17, "FIG_17_HIST_IID_pre.png",
     "Figura 17a. Histograma con curva normal — IID Pre-Test",
     "Distribución del IID en el pre-test (n = 43 días, gestión manual). "
     "Media 58.72%, SD 14.43%. La curva normal superpuesta evidencia la "
     "forma de la distribución para la prueba de Shapiro-Wilk."),

    (67, "FIG_67_HIST_IID_post.png",
     "Figura 17b. Histograma con curva normal — IID Post-Test",
     "Distribución del IID en el post-test (n = 24 días, con sistema). "
     "Media 97.52%, SD 5.68%. Nótese la concentración cercana al 100% "
     "con menor dispersión que el pre-test."),

    (18, "FIG_18_HIST_CHR_pre.png",
     "Figura 18a. Histograma con curva normal — CHR Pre-Test",
     "Distribución del CHR en el pre-test. Media 65.70%, SD 9.99%, n = 43 días."),

    (68, "FIG_68_HIST_CHR_post.png",
     "Figura 18b. Histograma con curva normal — CHR Post-Test",
     "Distribución del CHR en el post-test. Media 93.55%, SD 7.83%, n = 24 días. "
     "El Shapiro-Wilk arroja p = 0.100 > 0.05 → distribución normal."),

    (19, "FIG_19_HIST_TDE_pre.png",
     "Figura 19a. Histograma con curva normal — TDE Pre-Test",
     "Distribución del TDE en el pre-test. Media 51.74%, SD 15.60%, n = 43 días. "
     "El Shapiro-Wilk arroja p = 0.500 > 0.05 → distribución normal."),

    (69, "FIG_69_HIST_TDE_post.png",
     "Figura 19b. Histograma con curva normal — TDE Post-Test",
     "Distribución del TDE en el post-test. Media 93.35%, SD 8.07%, n = 24 días."),

    # ───── Sección 3: Curvas T-Student (Prueba de Hipótesis) ─────
    (20, "FIG_20_TSTUDENT_IID.png",
     "Figura 20. Prueba T-Student — IID",
     "El valor calculado Tc = 10.542 cae en la zona de rechazo (área roja), "
     "muy por encima del valor crítico T = 1.7139 (gl = 23, α = 0.05). "
     "→ Se rechaza H₀ y se acepta Hₐ: el aplicativo móvil incrementa "
     "significativamente la integridad de los datos registrados. "
     "Tamaño del efecto: d de Cohen = 2.152 (muy grande)."),

    (21, "FIG_21_TSTUDENT_CHR.png",
     "Figura 21. Prueba T-Student — CHR",
     "Tc = 9.168 > T crítico = 1.7139. Se rechaza H₀ y se acepta Hₐ: "
     "el aplicativo móvil incrementa significativamente el cumplimiento "
     "de la hoja de ruta. d de Cohen = 1.871 (muy grande)."),

    (22, "FIG_22_TSTUDENT_TDE.png",
     "Figura 22. Prueba T-Student — TDE",
     "Tc = 10.329 > T crítico = 1.7139. Se rechaza H₀ y se acepta Hₐ: "
     "el aplicativo móvil incrementa significativamente la tasa de "
     "disponibilidad de evidencias digitales. d de Cohen = 2.108 (muy grande)."),
]


def set_paragraph_border(p, color="999999", size=4):
    """Agrega borde inferior a un párrafo (para dividir secciones)."""
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_centered(doc, text, *, size=11, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_section_header(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
    set_paragraph_border(p, color="1F77B4", size=8)
    if subtitle:
        add_centered(doc, subtitle, size=11, italic=True, color=RGBColor(0x55, 0x55, 0x55))


def add_figure_page(doc, num, fname, titulo, descripcion, page_break=True):
    """Inserta una figura: heading + imagen centrada + pie + page break."""
    if page_break:
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    # Título de la figura
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Imagen
    img_path = FIG_DIR / fname
    if img_path.exists():
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(6)
        run = p_img.add_run()
        run.add_picture(str(img_path), width=Inches(6.2))
    else:
        add_centered(doc, f"[Imagen no encontrada: {fname}]",
                     size=10, italic=True, color=RGBColor(0xCC, 0x33, 0x33))

    # Descripción / pie de figura
    add_centered(doc, descripcion, size=10, italic=True,
                 color=RGBColor(0x44, 0x44, 0x44), space_after=4)
    add_centered(doc, "Fuente: Elaboración propia con datos del sistema EcoRoute.",
                 size=9, italic=True, color=RGBColor(0x88, 0x88, 0x88))


def add_toc_table(doc):
    """Índice de figuras del documento."""
    doc.add_heading("Índice de Figuras", level=1)
    rows = [
        ("Sección", "Figura", "Título", "Página"),
    ]

    # Sección 1
    rows.append(("3.1 Análisis Descriptivo", "Figura 14", "Comparativo IID", "3"))
    rows.append(("",                          "Figura 15", "Comparativo CHR", "4"))
    rows.append(("",                          "Figura 16", "Comparativo TDE", "5"))
    # Sección 2
    rows.append(("3.2 Análisis Inferencial",  "Figura 17a", "Histograma IID Pre", "6"))
    rows.append(("",                          "Figura 17b", "Histograma IID Post", "7"))
    rows.append(("",                          "Figura 18a", "Histograma CHR Pre", "8"))
    rows.append(("",                          "Figura 18b", "Histograma CHR Post", "9"))
    rows.append(("",                          "Figura 19a", "Histograma TDE Pre", "10"))
    rows.append(("",                          "Figura 19b", "Histograma TDE Post", "11"))
    # Sección 3
    rows.append(("3.3 Prueba de Hipótesis",   "Figura 20", "T-Student IID", "12"))
    rows.append(("",                          "Figura 21", "T-Student CHR", "13"))
    rows.append(("",                          "Figura 22", "T-Student TDE", "14"))

    tbl = doc.add_table(rows=len(rows), cols=4)
    tbl.style = 'Light Grid Accent 1'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            if r_idx == 0:
                run.bold = True
                # Background azul
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), '1F77B4')
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(9)
    doc.add_paragraph()


def main():
    print("Generando docx con los 12 gráficos del Cap. III...")
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ─── Página 1: portada ───
    add_centered(doc, "GRÁFICOS DEL CAPÍTULO III", size=22, bold=True,
                 color=RGBColor(0x1F, 0x29, 0x37), space_after=4)
    add_centered(doc, "Resultados Estadísticos", size=16, bold=True,
                 color=RGBColor(0x1F, 0x77, 0xB4), space_after=20)
    add_centered(doc, "Tesis: Aplicativo móvil para la gestión administrativa",
                 size=13, italic=True, color=RGBColor(0x44, 0x44, 0x44), space_after=2)
    add_centered(doc, "Grupo Micotrans S.A.C. — Puente Piedra, Lima 2026",
                 size=11, italic=True, color=RGBColor(0x55, 0x55, 0x55), space_after=30)
    add_centered(doc, "Autor: Campos Vargas, Kevin Stip",
                 size=11, color=RGBColor(0x44, 0x44, 0x44), space_after=2)
    add_centered(doc, "ORCID: 0000-0002-6087-3626",
                 size=10, italic=True, color=RGBColor(0x88, 0x88, 0x88), space_after=40)

    # Resumen estadístico en caja
    add_centered(doc, "Resumen Estadístico (n = 24 días pareados, T crítico = 1.7139)",
                 size=12, bold=True, color=RGBColor(0x1F, 0x77, 0xB4))
    summary = doc.add_table(rows=4, cols=5)
    summary.style = 'Light Grid Accent 1'
    headers = ["Indicador", "Pre M (%)", "Post M (%)", "Tc", "Decisión"]
    data = [
        ("IID", "58.72", "97.52", "10.542", "Acepta Hₐ ✓"),
        ("CHR", "65.70", "93.55",  "9.168", "Acepta Hₐ ✓"),
        ("TDE", "51.74", "93.35", "10.329", "Acepta Hₐ ✓"),
    ]
    for c, h in enumerate(headers):
        cell = summary.cell(0, c)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '1F77B4')
        tcPr.append(shd)
    for r_idx, row in enumerate(data, start=1):
        for c_idx, val in enumerate(row):
            summary.cell(r_idx, c_idx).text = val
            summary.cell(r_idx, c_idx).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ─── Página 2: Índice de figuras ───
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    add_toc_table(doc)

    # ─── Páginas 3+ : las 12 figuras ───
    # Cada figura en su propia página con encabezado de sección
    section_seen = set()
    for num, fname, titulo, descripcion in FIGURAS:
        # Determinar sección
        if num in (14, 15, 16):
            section = ("3.1 Análisis Descriptivo",
                       "Comparativos Pre-Test vs Post-Test (estadísticos descriptivos)")
        elif num in (17, 18, 19, 67, 68, 69):
            section = ("3.2 Análisis Inferencial",
                       "Histogramas con curva normal — Prueba de Shapiro-Wilk")
        else:
            section = ("3.3 Prueba de Hipótesis",
                       "Curvas T-Student con zona de rechazo")

        # Imprimir cabecera de sección sólo en la primera figura de cada bloque
        if section[0] not in section_seen:
            section_seen.add(section[0])
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            add_section_header(doc, section[0], section[1])
            add_figure_page(doc, num, fname, titulo, descripcion, page_break=False)
        else:
            add_figure_page(doc, num, fname, titulo, descripcion, page_break=True)

    # ─── Última página: conclusión ───
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    add_section_header(doc, "Conclusión",
                       "Contrastación de la Hipótesis General")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(
        "Los resultados estadísticos demuestran que el aplicativo móvil EcoRoute, "
        "desplegado en la empresa Grupo Micotrans S.A.C., incrementa "
        "significativamente los tres indicadores que componen la variable "
        "dependiente \"gestión administrativa\": IID, CHR y TDE. "
        "Las tres hipótesis específicas se contrastaron mediante la prueba T de "
        "Student para muestras pareadas, obteniendo en los tres casos un valor "
        "calculado superior al valor crítico de 1.7139 (gl = 23, α = 0.05), con "
        "magnitudes de efecto muy grandes (d de Cohen > 1.85) y nivel de "
        "significancia bilateral p < 0.001. Por lo tanto, se cumple la "
        "hipótesis general del estudio."
    )
    run.font.size = Pt(11)

    doc.save(str(OUT))
    print(f"\n[OK] Generado: {OUT}")
    print(f"     Tamaño: {OUT.stat().st_size // 1024} KB")
    print(f"     Páginas estimadas: ~16 (portada + índice + 12 figuras + conclusión)")


if __name__ == "__main__":
    main()
