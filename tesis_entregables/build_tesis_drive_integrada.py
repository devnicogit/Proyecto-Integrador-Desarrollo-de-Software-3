"""
build_tesis_drive_integrada.py
==============================

Toma el documento REAL más actualizado de Kevin (descargado del Drive como .docx)
y le INSERTA / COMPLETA las secciones que están vacías, sin destruir lo que Kevin
ya tiene escrito (carátula, Cap I Introducción, Cap II Metodología, 51 imágenes,
Anexo 8 con sus CAPÍTULO IV Programación y V Pruebas Calidad, etc.).

Operaciones que realiza:
  1. Convierte el placeholder "Resultados" (vacío) en Heading 1 "III. Resultados"
     e inserta el contenido completo de Capitulo_3_Resultados.docx.
  2. Lo mismo para Discusión / Conclusiones / Recomendaciones (Cap IV-VI).
  3. Inserta el cuestionario UTAUT y las 3 fichas post-test después de las
     fichas pre-test ya existentes en el Anexo 2.
  4. Inserta el contenido pre-llenado en los Anexos 3, 4, 5 (hoy sólo tienen
     el título).
  5. Agrega "Anexo 9 - Figuras del sistema" al final con las 31 figuras
     (10 web admin + 21 móvil end-to-end) del Anexo_Figuras_Capturas.docx.

Entrada:  /c/tmp/kevin_drive_current.docx (descargado del Drive)
Salida:   tesis_entregables/docx/Tesis_Drive_Integrada_VFinal.docx
"""
from __future__ import annotations
import re
import shutil
import subprocess
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

HERE     = Path(__file__).resolve().parent
DOCX_DIR = HERE / "docx"
SRC_DOC  = Path(r"C:\tmp\kevin_drive_current.docx")
OUT      = DOCX_DIR / "Tesis_Drive_Integrada_VFinal.docx"

DOC_URL_ID = "1voac3fBjNJLK2rVc-uSCpvO99K5olD8x"  # del link compartido


# Relationship types que NO deben copiarse de los archivos fuente
# (los del target deben quedar intactos)
SKIP_RELTYPES = {
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer",
    "http://schemas.microsoft.com/office/2011/relationships/people",
    "http://schemas.microsoft.com/office/2011/relationships/commentsExtended",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml",
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/stylesWithEffects",
}


def download_drive_doc() -> None:
    """Descarga el doc del Drive (si no está ya en cache)."""
    SRC_DOC.parent.mkdir(parents=True, exist_ok=True)
    if SRC_DOC.exists() and SRC_DOC.stat().st_size > 100_000:
        print(f"  Cache: {SRC_DOC} ({SRC_DOC.stat().st_size//1024} KB)")
        return
    url = f"https://docs.google.com/document/d/{DOC_URL_ID}/export?format=docx"
    subprocess.run(["curl", "-sL", "-o", str(SRC_DOC), url], check=True)
    if SRC_DOC.stat().st_size < 100_000:
        raise RuntimeError(
            f"Descarga falló o el doc no es público. Tamaño: {SRC_DOC.stat().st_size}\n"
            f"Compartí el doc como 'cualquiera con el link puede ver' y reintentá."
        )
    print(f"  Descargado: {SRC_DOC} ({SRC_DOC.stat().st_size//1024} KB)")


# --------------------------------------------------------------------------- #
# Helpers de inserción de elementos XML
# --------------------------------------------------------------------------- #

def _map_rels(target_part, src_part) -> dict[str, str]:
    """Devuelve un dict {old_rid_in_src: new_rid_in_target} para los rels
    relevantes (excluye estilos/numerado/etc.)."""
    rid_map = {}
    for rel_id, rel in src_part.rels.items():
        if rel.reltype in SKIP_RELTYPES:
            continue
        if rel.is_external:
            new_rid = target_part.relate_to(rel.target_ref, rel.reltype, is_external=True)
        else:
            new_rid = target_part.relate_to(rel.target_part, rel.reltype)
        rid_map[rel_id] = new_rid
    return rid_map


def _fix_rids(element, rid_map: dict[str, str]) -> None:
    """Remapea los rIds de imágenes e hyperlinks dentro del elemento clonado."""
    for blip in element.iter(qn("a:blip")):
        old = blip.get(qn("r:embed"))
        if old in rid_map:
            blip.set(qn("r:embed"), rid_map[old])
    for hyper in element.iter(qn("w:hyperlink")):
        old = hyper.get(qn("r:id"))
        if old in rid_map:
            hyper.set(qn("r:id"), rid_map[old])
    # v:imagedata (legacy Word) — namespace v: no siempre registrado, omitir


def insert_docx_after(target_doc: Document, target_paragraph, source_path: Path,
                       skip_first_heading: bool = True) -> int:
    """Inserta el cuerpo del source DESPUÉS del target_paragraph.

    Si skip_first_heading=True, omite el primer Heading 1 del source
    (porque el target ya tiene el título de la sección).

    Devuelve la cantidad de elementos insertados.
    """
    if not source_path.exists():
        print(f"  WARN: falta {source_path.name}")
        return 0
    src = Document(str(source_path))
    rid_map = _map_rels(target_doc.part, src.part)
    cursor = target_paragraph._element
    inserted = 0
    skipped_first = False
    for element in src.element.body:
        if element.tag == qn("w:sectPr"):
            continue
        # Omitir primer Heading 1 si pide
        if skip_first_heading and not skipped_first:
            is_h1 = False
            for pStyle in element.iter(qn("w:pStyle")):
                if pStyle.get(qn("w:val")) in ("Heading1", "Ttulo1", "Title"):
                    is_h1 = True
                    break
            if is_h1:
                skipped_first = True
                continue
            # También saltar párrafos vacíos previos
            if element.tag == qn("w:p") and not element.findall(qn("w:r")):
                continue
        new_element = deepcopy(element)
        _fix_rids(new_element, rid_map)
        cursor.addnext(new_element)
        cursor = new_element
        inserted += 1
    return inserted


def append_docx_at_end(target_doc: Document, source_path: Path,
                         page_break_before: bool = True) -> int:
    """Agrega el cuerpo del source AL FINAL del target (antes del sectPr)."""
    if not source_path.exists():
        print(f"  WARN: falta {source_path.name}")
        return 0
    src = Document(str(source_path))
    rid_map = _map_rels(target_doc.part, src.part)
    body = target_doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if page_break_before:
        p = target_doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
    inserted = 0
    for element in src.element.body:
        if element.tag == qn("w:sectPr"):
            continue
        new_element = deepcopy(element)
        _fix_rids(new_element, rid_map)
        if sect_pr is not None:
            sect_pr.addprevious(new_element)
        else:
            body.append(new_element)
        inserted += 1
    return inserted


def find_paragraph_by_exact_text(doc: Document, text: str, start_from: int = 0) -> int | None:
    """Encuentra el índice del primer párrafo cuyo texto.strip() == text."""
    for i, p in enumerate(doc.paragraphs):
        if i < start_from:
            continue
        if p.text.strip() == text:
            return i
    return None


def promote_to_heading1(paragraph, new_text: str) -> None:
    """Convierte un párrafo "normal" en Heading 1 con el texto dado."""
    # Cambiar estilo
    paragraph.style = paragraph.part.document.styles["Heading 1"]
    # Cambiar texto: limpiar runs existentes y poner uno nuevo
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    run = paragraph.add_run(new_text)
    run.bold = True
    run.font.size = Pt(16)


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #

def main() -> None:
    print("== build_tesis_drive_integrada.py ==")
    download_drive_doc()

    # Copia el doc del Drive como base de trabajo
    OUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(str(SRC_DOC), str(OUT))
    doc = Document(str(OUT))
    print(f"  Base cargada: {len(doc.paragraphs)} párrafos, "
          f"{sum(1 for r in doc.part.rels.values() if 'image' in r.reltype)} imágenes")

    # ----- 1. Cap III Resultados -----
    idx = find_paragraph_by_exact_text(doc, "Resultados")
    if idx is None:
        print("  WARN: no encuentro placeholder 'Resultados'")
    else:
        print(f"  [{idx}] 'Resultados' -> promote a H1 'III. Resultados' + insertar contenido")
        promote_to_heading1(doc.paragraphs[idx], "III. Resultados")
        n = insert_docx_after(doc, doc.paragraphs[idx],
                                DOCX_DIR / "Capitulo_3_Resultados.docx")
        print(f"     -> {n} elementos insertados")

    # ----- 2. Cap IV Discusión / V Conclusiones / VI Recomendaciones -----
    # Estos 3 vienen TODOS en Capitulo_4_5_6_Discusion_Conclusiones_Recomendaciones.docx
    # Lo insertamos después de "Discusión" y los placeholders V y VI quedarán como duplicados.
    # Solución: insertarlo después de "Discusión" y luego BORRAR los placeholders V y VI.
    idx_iv = find_paragraph_by_exact_text(doc, "Discusión")
    idx_v  = find_paragraph_by_exact_text(doc, "Conclusiones")
    idx_vi = find_paragraph_by_exact_text(doc, "Recomendaciones")
    if idx_iv is not None:
        print(f"  [{idx_iv}] 'Discusión' -> promote a H1 'IV. Discusión' + insertar Cap IV-V-VI")
        promote_to_heading1(doc.paragraphs[idx_iv], "IV. Discusión")
        n = insert_docx_after(doc, doc.paragraphs[idx_iv],
                               DOCX_DIR / "Capitulo_4_5_6_Discusion_Conclusiones_Recomendaciones.docx")
        print(f"     -> {n} elementos insertados (incluye IV, V, VI)")
        # Ahora borrar los placeholders V y VI residuales del Drive original.
        # Tras insertar Cap IV-V-VI quedan dos párrafos "Normal" con texto exacto
        # "Conclusiones" y "Recomendaciones" (los headings reales del Cap V/VI ya
        # quedaron como Heading 1 al ser insertados desde el .docx). Los borramos.
        for ph_text in ("Conclusiones", "Recomendaciones"):
            for p in list(doc.paragraphs):
                if p.text.strip() == ph_text and p.style.name == "Normal":
                    parent = p._element.getparent()
                    parent.remove(p._element)
                    print(f"     -> placeholder '{ph_text}' eliminado")
                    break

    # ----- 3. Anexo 2: agregar UTAUT + fichas post-test -----
    # En el Drive ya hay fichas PRE-TEST. Agregar al final del Anexo 2:
    # - Anexo_2_Cuestionario_UTAUT.docx (instrumento)
    # - Las fichas post-test podemos integrarlas también, pero requieren PDF embebido.
    #   Por simplicidad agregamos el Anexo_2_Consolidado (que ya tiene UTAUT + 6 fichas).
    # Buscamos el inicio de Anexo 3 para insertar ANTES.
    idx_anexo3 = find_paragraph_by_exact_text(doc,
        "Anexo 3. Ficha de validación de contenido para un instrumento")
    if idx_anexo3 is not None:
        target = doc.paragraphs[idx_anexo3 - 1] if idx_anexo3 > 0 else doc.paragraphs[idx_anexo3]
        print(f"  [{idx_anexo3}] insertar UTAUT antes de Anexo 3")
        # Heading separador
        h = doc.add_paragraph("Anexo 2 (continuación). Cuestionario UTAUT y fichas post-test")
        # Lo movemos a la posición correcta
        target._element.addnext(deepcopy(h._element))
        h._element.getparent().remove(h._element)
        # Insertar el contenido del cuestionario UTAUT
        n = insert_docx_after(doc, doc.paragraphs[idx_anexo3 - 1],
                               DOCX_DIR / "Anexo_2_Cuestionario_UTAUT.docx",
                               skip_first_heading=False)
        print(f"     -> UTAUT: {n} elementos insertados")

    # ----- 4. Anexo 3: validación juicio expertos -----
    idx_anexo3 = find_paragraph_by_exact_text(doc,
        "Anexo 3. Ficha de validación de contenido para un instrumento")
    if idx_anexo3 is not None:
        print(f"  [{idx_anexo3}] insertar contenido Anexo 3")
        n = insert_docx_after(doc, doc.paragraphs[idx_anexo3],
                               DOCX_DIR / "Anexo_3_Validacion_Juicio_Expertos_PreLlenado.docx",
                               skip_first_heading=True)
        print(f"     -> {n} elementos insertados")

    # ----- 5. Anexo 4: solicitud autorización -----
    idx_anexo4 = find_paragraph_by_exact_text(doc,
        "Anexo 4. Solicitud de autorización para realizar la investigación en una institución")
    if idx_anexo4 is not None:
        print(f"  [{idx_anexo4}] insertar contenido Anexo 4")
        n = insert_docx_after(doc, doc.paragraphs[idx_anexo4],
                               DOCX_DIR / "Anexo_4_Solicitud_Autorizacion.docx",
                               skip_first_heading=True)
        print(f"     -> {n} elementos insertados")

    # ----- 6. Anexo 5: autorización MICOTRANS -----
    idx_anexo5 = find_paragraph_by_exact_text(doc,
        "Anexo 5. Autorización de uso de información de la empresa")
    if idx_anexo5 is not None:
        print(f"  [{idx_anexo5}] insertar contenido Anexo 5")
        n = insert_docx_after(doc, doc.paragraphs[idx_anexo5],
                               DOCX_DIR / "Anexo_5_Autorizacion_MICOTRANS.docx",
                               skip_first_heading=True)
        print(f"     -> {n} elementos insertados")

    # ----- 7. Anexo Figuras del sistema (NUEVO, al final) -----
    print("  + Anexo Figuras del sistema (31 figuras web + móvil)")
    n = append_docx_at_end(doc, DOCX_DIR / "Anexo_Figuras_Capturas.docx",
                             page_break_before=True)
    print(f"     -> {n} elementos insertados")

    # ----- Guardar -----
    doc.save(str(OUT))
    print(f"\nOK -> {OUT}")
    print(f"   Tamaño: {OUT.stat().st_size // 1024} KB")
    # Stats finales
    doc2 = Document(str(OUT))
    print(f"   Párrafos finales: {len(doc2.paragraphs)}")
    print(f"   Imágenes finales: {sum(1 for r in doc2.part.rels.values() if 'image' in r.reltype)}")


if __name__ == "__main__":
    main()
