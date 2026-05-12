"""
Inserta las capturas PNG reales en los documentos .docx existentes
en los lugares marcados por las referencias 'Figura N°XX'.

Estrategia: re-escribe los .docx desde los .md originales pero detecta
las menciones 'Figura XX' y agrega la imagen correspondiente debajo del
párrafo que la menciona.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding='utf-8')

ROOT = Path(__file__).parent
DOCX = ROOT / "docx"
CAP = ROOT / "figuras_capturas"

# Mapeo de menciones a archivos de figura
FIGURE_MAP = {
    "Figura 20": ("F20_login.png", "Figura 20: Pantalla de Login del Panel Administrativo EcoRoute."),
    "Figura 21": ("F21_home.png", "Figura 21: Pantalla Home / Bienvenida al sistema."),
    "Figura 22": ("F22_dashboard.png", "Figura 22: Panel de Control Logístico con KPIs operativos."),
    "Figura 23": ("F23_pedidos.png", "Figura 23: Gestión de Pedidos mostrando los datos reales de MICOTRANS (guías GR-1001 a GR-1010)."),
    "Figura 24": ("F24_rutas.png", "Figura 24: Planificación y Seguimiento de Rutas con mapa y polilínea."),
    "Figura 25": ("F25_conductores.png", "Figura 25: Gestión de los 5 conductores asignados a MICOTRANS."),
    "Figura 26": ("F26_vehiculos.png", "Figura 26: Gestión de la flota AFT-101 a AFT-105 de MICOTRANS."),
    "Figura 39": ("F39_dashboard_top.png", "Figura 39: Dashboard en ejecución mostrando los pedidos del día."),
    "Figura 47": ("F47_kpi_dashboard_post.png", "Figura 47: Dashboard de KPIs de Gestión Administrativa (Tesis) — vista Post-Test mostrando IID 97.8%, CHR 93.5%, TDE 94.1%."),
}

# Figuras adicionales que se anexan al final
EXTRA_FIGURES = [
    ("F47b_kpi_dashboard_pre.png", "Figura 47-B: Dashboard de KPIs vista Pre-Test mostrando los valores reales del CSV de MICOTRANS: IID 60.0%, CHR 67.3%, TDE 51.3%."),
]


def add_image_after_paragraph(doc, paragraph, image_path, caption):
    """Inserta una imagen + caption como párrafos nuevos justo después del párrafo dado."""
    # Crear nuevo paragraph para la imagen
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(6.0))
    except Exception as e:
        run.add_text(f"[ERROR cargando imagen {image_path}: {e}]")

    # Mover el nuevo párrafo justo después del párrafo de referencia
    p_after = paragraph._p
    new_p = new_para._p
    p_after.addnext(new_p)

    # Caption
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    new_p.addnext(cap_para._p)


def insert_figures(docx_path):
    """Abre un docx y para cada párrafo que contenga 'Figura XX' inserta la imagen."""
    doc = Document(str(docx_path))
    inserted = []

    # Necesitamos iterar pero también modificar - hacemos snapshot primero
    paragraphs = list(doc.paragraphs)

    for para in paragraphs:
        text = para.text
        for fig_key, (img_file, caption) in FIGURE_MAP.items():
            # Match exacto "Figura 47" pero NO "Figura 47B"
            if re.search(rf"\b{fig_key}\b", text) and fig_key not in inserted:
                img_path = CAP / img_file
                if img_path.exists():
                    add_image_after_paragraph(doc, para, img_path, caption)
                    inserted.append(fig_key)

    doc.save(str(docx_path))
    return inserted


def main():
    # Documentos donde tiene sentido insertar figuras
    targets = [
        "Anexo_8_Metodologia.docx",
        "Capitulo_3_Resultados.docx",
    ]

    for name in targets:
        path = DOCX / name
        if not path.exists():
            print(f"SKIP no existe: {name}")
            continue
        print(f"\n=== {name} ===")
        inserted = insert_figures(path)
        for fig in inserted:
            print(f"  OK insertada: {fig}")
        if not inserted:
            print(f"  (sin figuras detectadas en el texto)")

    # Crear documento Anexo de Figuras adicional con todas las capturas
    print("\n=== Creando Anexo_Figuras_Capturas.docx ===")
    create_figuras_anexo()


def create_figuras_anexo():
    doc = Document()
    title = doc.add_heading('Anexo Visual - Capturas del Sistema EcoRoute', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Este anexo contiene las capturas reales del sistema EcoRoute en operación, "
        "tomadas durante el periodo del post-test (mayo 2026) con la base de datos "
        "cargada con los registros reales de Grupo Micotrans S.A.C."
    )
    intro_run.italic = True
    doc.add_paragraph()

    # Todas las figuras en orden
    all_figs = list(FIGURE_MAP.values()) + EXTRA_FIGURES
    for img_file, caption in all_figs:
        img_path = CAP / img_file
        if not img_path.exists():
            continue
        # Caption como heading menor
        h = doc.add_heading(caption.split(':')[0], level=3)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        try:
            run.add_picture(str(img_path), width=Inches(6.0))
        except Exception as e:
            run.add_text(f"[ERROR cargando {img_file}: {e}]")

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()  # espacio

    out = DOCX / "Anexo_Figuras_Capturas.docx"
    doc.save(str(out))
    print(f"  OK creado: {out.name} ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
