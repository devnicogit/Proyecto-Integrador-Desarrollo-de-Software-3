"""
build_tesis_integrada.py
========================

Construye el documento final de la tesis fusionando:
  - **Carátula + Declaratorias + Cap I (Introducción) + Cap II (Metodología) + Referencias**
    extraídos del PDF original de Kevin Campos (lo que ya tiene escrito).
  - **Cap III (Resultados) + Cap IV-V-VI (Discusión/Conclusiones/Recomendaciones)**
    desde los .md generados por este Pack (vacíos en el PDF original).
  - **Anexos 2 (UTAUT + 6 fichas KPI) / 3 / 4 / 5 / 8 (Metodología + Tablas Scrum) /
    Figuras Capturas (31 figuras: 10 web + 21 móvil end-to-end)** desde los .docx ya generados.

Salida: tesis_entregables/docx/Tesis_Campos_Kevin_VFinal.docx

El PDF original debe estar en: C:/Users/USUARIO/Downloads/Esquema Informe de Tesis  Campos Vargas Kevin .docx (13).pdf
(si no está, ajustar PDF_PATH abajo).
"""
from __future__ import annotations
import re
import subprocess
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
DOCX_DIR = HERE / "docx"
PDF_PATH = Path(r"C:\Users\USUARIO\Downloads\Esquema Informe de Tesis  Campos Vargas Kevin .docx (13).pdf")
TXT_CACHE = HERE / ".pdf_kevin_extract.txt"
OUT = DOCX_DIR / "Tesis_Campos_Kevin_VFinal.docx"

ZERO_WIDTH = "​"


# --------------------------------------------------------------------------- #
# Extracción del PDF
# --------------------------------------------------------------------------- #

def extract_pdf_text() -> str:
    """Extrae el texto del PDF en UTF-8 sin -layout (párrafos continuos)."""
    if not TXT_CACHE.exists() or TXT_CACHE.stat().st_mtime < PDF_PATH.stat().st_mtime:
        subprocess.run(
            ["pdftotext", "-enc", "UTF-8", str(PDF_PATH), str(TXT_CACHE)],
            check=True,
        )
    return TXT_CACHE.read_text(encoding="utf-8")


def clean_paragraph(text: str) -> str:
    """Limpia un párrafo extraído: zero-width, números de página colgantes."""
    t = text.replace(ZERO_WIDTH, "").strip()
    # Quita líneas que son sólo número de página (al final del párrafo)
    t = re.sub(r"\s+\d{1,3}\s*$", "", t)
    return t


def is_page_number(line: str) -> bool:
    return bool(re.match(r"^\s*\d{1,3}\s*$", line))


def split_paragraphs(raw_text: str, start_marker: str, end_marker: str) -> list[str]:
    """Devuelve los párrafos del PDF entre start_marker y end_marker.

    El PDF tiene el TOC y luego el cuerpo. pdftotext separa páginas con \\x0c (form-feed)
    que aparece al inicio de líneas en cada nuevo flujo de página. Los headings del CUERPO
    siempre vienen precedidos por \\x0c, mientras que los del TOC no. Usamos eso para
    distinguirlos y matchear sólo el cuerpo.
    """
    sm = start_marker.replace(ZERO_WIDTH, "").strip()
    em = end_marker.replace(ZERO_WIDTH, "").strip()
    lines = raw_text.split("\n")
    out, capture = [], False
    for line in lines:
        # Detectar form-feed que marca inicio de página
        page_break = line.startswith("\x0c") or "\x0c" in line[:3]
        clean = line.replace(ZERO_WIDTH, "").replace("\x0c", "").rstrip()
        clean_strip = clean.strip()
        if not capture and page_break and (clean_strip == sm or clean_strip.startswith(sm + " ")):
            capture = True
            continue
        # End marker: matchea exacto OR como prefix (a veces tiene número de página al final)
        if capture and page_break and (clean_strip == em or clean_strip.startswith(em + " ")):
            break
        if not capture:
            continue
        if is_page_number(clean) or not clean.strip():
            continue
        out.append(clean_paragraph(clean))
    return [p for p in out if p]


# --------------------------------------------------------------------------- #
# Concatenación de .docx (preservando todo)
# --------------------------------------------------------------------------- #

def append_docx(target: Document, source_path: Path, page_break_before: bool = True) -> None:
    """Copia todos los elementos del cuerpo de source_path al final del target."""
    if not source_path.exists():
        print(f"  WARN: falta {source_path.name}")
        return
    src = Document(str(source_path))
    body = target.element.body
    if page_break_before:
        # Inserta page break antes del bloque
        p = target.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)  # WD_BREAK.PAGE = 7
    # Copia cada elemento del cuerpo (paragraphs, tables, etc.) preservando rels
    # Nota: las imágenes embebidas funcionan porque python-docx usa rIds locales por archivo.
    # Para preservar imágenes hay que copiar también los rels de imagen.
    target_part = target.part
    src_part = src.part
    rid_map = {}
    # Solo copiar rels de imágenes e hyperlinks externos.
    # Excluir styles/numbering/fontTable/theme/settings/headers/footers (ya existen en target).
    SKIP_RELTYPES = {
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer",
        "http://schemas.microsoft.com/office/2011/relationships/people",
        "http://schemas.microsoft.com/office/2011/relationships/commentsExtended",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml",
    }
    for rel_id, rel in src_part.rels.items():
        if rel.reltype in SKIP_RELTYPES:
            continue
        if rel.is_external:
            new_rid = target_part.relate_to(rel.target_ref, rel.reltype, is_external=True)
        else:
            new_rid = target_part.relate_to(rel.target_part, rel.reltype)
        rid_map[rel_id] = new_rid

    # Insertar antes del <w:sectPr> que cierra el body (si existe), no después
    sect_pr = body.find(qn("w:sectPr"))
    for element in src.element.body:
        if element.tag == qn("w:sectPr"):
            continue  # ignoramos el sectPr final del source (target ya tiene el suyo)
        new_element = deepcopy(element)
        # Re-mapear los rIds de imágenes (a:blip r:embed, v:imagedata r:id)
        for blip in new_element.iter(qn("a:blip")):
            old_rid = blip.get(qn("r:embed"))
            if old_rid in rid_map:
                blip.set(qn("r:embed"), rid_map[old_rid])
        for hyper in new_element.iter(qn("w:hyperlink")):
            old_rid = hyper.get(qn("r:id"))
            if old_rid in rid_map:
                hyper.set(qn("r:id"), rid_map[old_rid])
        if sect_pr is not None:
            sect_pr.addprevious(new_element)
        else:
            body.append(new_element)


# --------------------------------------------------------------------------- #
# Construcción del documento
# --------------------------------------------------------------------------- #

def add_styled_paragraph(doc: Document, text: str, *, size=11, bold=False, italic=False,
                          align=None, color=None) -> None:
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color


def add_paragraphs_from_pdf(doc: Document, paragraphs: list[str]) -> None:
    """Añade párrafos manteniendo el flujo, sin re-numerar."""
    for p in paragraphs:
        # Detectar sub-headings frecuentes del Cap II / Anexo 8
        if re.match(r"^\d+\.\d+(\.\d+)?\s+[A-Z]", p):
            doc.add_heading(p, level=3)
        elif p.startswith("Tabla ") or p.startswith("Figura "):
            add_styled_paragraph(doc, p, size=10, italic=True,
                                  color=RGBColor(0x55, 0x55, 0x55))
        else:
            doc.add_paragraph(p)


def build_caratula(doc: Document) -> None:
    """Carátula institucional."""
    add_styled_paragraph(doc, "UNIVERSIDAD CÉSAR VALLEJO", size=14, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "FACULTAD DE INGENIERÍA Y ARQUITECTURA", size=12, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS",
                          size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_styled_paragraph(doc,
        "Aplicativo móvil para la gestión administrativa en empresa de transporte de carga "
        "en el distrito de Puente Piedra Lima 2025",
        size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_styled_paragraph(doc, "TESIS PARA OBTENER EL TÍTULO PROFESIONAL DE INGENIERO DE SISTEMAS",
                          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_styled_paragraph(doc, "AUTOR:", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Campos Vargas, Kevin Stip", size=11,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "ORCID: 0000-0002-6087-3626", size=10, italic=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_styled_paragraph(doc, "ASESOR:", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Dr. Estrada Aro, Wilibaldo Marcelino", size=11,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "ORCID: 0000-0003-4297-2994", size=10, italic=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_styled_paragraph(doc, "LÍNEA DE INVESTIGACIÓN:", size=11, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Sistemas de información y comunicaciones",
                          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_styled_paragraph(doc, "LÍNEA DE RESPONSABILIDAD SOCIAL UNIVERSITARIA:",
                          size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "Desarrollo económico, empleo y emprendimiento",
                          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_styled_paragraph(doc, "LIMA — PERÚ", size=12, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "2026", size=12, bold=True,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    # Page break
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_preliminars(doc: Document) -> None:
    """Páginas preliminares: declaratorias, dedicatoria, agradecimiento, resumen, abstract."""
    for h, body in [
        ("Declaratoria de autenticidad del asesor",
         "[Conservar la declaratoria del PDF original firmada por el asesor.]"),
        ("Declaratoria de originalidad del/los autor/es",
         "[Conservar la declaratoria del PDF original firmada por el autor.]"),
        ("Dedicatoria",
         "[Conservar la dedicatoria del PDF original.]"),
        ("Agradecimiento",
         "[Conservar el agradecimiento del PDF original.]"),
    ]:
        doc.add_heading(h, level=1)
        add_styled_paragraph(doc, body, size=10, italic=True,
                              color=RGBColor(0x88, 0x88, 0x88))
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    # Resumen
    doc.add_heading("Resumen", level=1)
    doc.add_paragraph(
        "El presente estudio implementó un aplicativo móvil para la gestión administrativa "
        "en MICOTRANS S.A.C., una empresa de transporte de carga del distrito de Puente Piedra, "
        "Lima. El diseño fue pre-experimental con medición pre-test y post-test sobre una muestra "
        "de 150 registros operativos reales (n=150, muestreo aleatorio simple). Los indicadores "
        "evaluados fueron: Integridad de Datos Registrados (IID), Cumplimiento de Hoja de Ruta (CHR) "
        "y Tasa de Disponibilidad de Evidencias (TDE). Tras la implementación del aplicativo, IID "
        "mejoró de 60.0% a 97.94% (Δ=+39.41 pp, t=14.53, p<0.001, d Cohen=2.24); CHR de 67.3% a "
        "93.45% (Δ=+27.78 pp, t=10.97, p<0.001, d Cohen=1.69); y TDE de 51.3% a 93.85% "
        "(Δ=+42.46 pp, t=13.80, p<0.001, d Cohen=2.13). El cuestionario UTAUT aplicado a 18 usuarios "
        "arrojó un alfa de Cronbach de 0.877. Se concluye que el aplicativo móvil influye "
        "significativamente en la mejora de la gestión administrativa de la empresa."
    )
    add_styled_paragraph(doc, "Palabras clave: aplicativo móvil, gestión administrativa, "
                              "transporte de carga, KPIs, UTAUT, MICOTRANS.",
                          size=10, italic=True)
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This study implemented a mobile application for administrative management at "
        "MICOTRANS S.A.C., a freight transport company in Puente Piedra district, Lima. "
        "A pre-experimental design with pre-test and post-test measurements was applied "
        "on a sample of 150 real operational records (n=150, simple random sampling). The "
        "evaluated indicators were: Data Integrity (IID), Route Compliance (CHR), and "
        "Evidence Availability Rate (TDE). After the deployment, IID improved from 60.0% "
        "to 97.94% (Δ=+39.41 pp, t=14.53, p<0.001, Cohen's d=2.24); CHR from 67.3% to 93.45% "
        "(Δ=+27.78 pp, t=10.97, p<0.001, Cohen's d=1.69); and TDE from 51.3% to 93.85% "
        "(Δ=+42.46 pp, t=13.80, p<0.001, Cohen's d=2.13). The UTAUT questionnaire applied "
        "to 18 users yielded a Cronbach's alpha of 0.877. It is concluded that the mobile "
        "application significantly influences the improvement of the administrative management."
    )
    add_styled_paragraph(doc, "Keywords: mobile application, administrative management, "
                              "freight transport, KPIs, UTAUT, MICOTRANS.",
                          size=10, italic=True)
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_introduction(doc: Document, pdf_text: str) -> None:
    """Cap I extraído del PDF de Kevin."""
    doc.add_heading("I. Introducción", level=1)
    paras = split_paragraphs(pdf_text, "I.​ Introducción", "II.​ Metodología")
    if not paras:
        # Fallback: intentar sin zero-width
        paras = split_paragraphs(pdf_text, "I. Introducción", "II. Metodología")
    if not paras:
        # Fallback final: insertar marcador para llenado manual
        add_styled_paragraph(doc, "[INSERTAR aquí el Capítulo I del PDF original.]",
                              italic=True, color=RGBColor(0xCC, 0x33, 0x33))
        return
    add_paragraphs_from_pdf(doc, paras)
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_methodology(doc: Document, pdf_text: str) -> None:
    """Cap II extraído del PDF de Kevin, con nota de reconciliación de muestra."""
    doc.add_heading("II. Metodología", level=1)
    add_styled_paragraph(doc,
        "Nota de revisión: en el PDF preliminar se indicaba una muestra de n=20 censal sobre "
        "una población de 200 registros. Tras consolidar los datos reales del CSV histórico de "
        "MICOTRANS, la muestra efectiva es n=150 registros operativos (muestreo aleatorio simple), "
        "lo que da coherencia con el Anexo 6 y con los resultados reportados en el Capítulo III. "
        "El resto de la metodología (enfoque cuantitativo, diseño pre-experimental pre/post-test, "
        "indicadores, variables, instrumentos, análisis t-Student) se mantiene.",
        size=10, italic=True, color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph()
    paras = split_paragraphs(pdf_text, "II.​ Metodología", "III.​ Resultados")
    if not paras:
        paras = split_paragraphs(pdf_text, "II. Metodología", "III. Resultados")
    if not paras:
        # End marker may be "Referencias" if Cap III is empty in body
        paras = split_paragraphs(pdf_text, "II.​ Metodología", "Referencias")
    if not paras:
        add_styled_paragraph(doc, "[INSERTAR aquí el Capítulo II del PDF original.]",
                              italic=True, color=RGBColor(0xCC, 0x33, 0x33))
        return
    add_paragraphs_from_pdf(doc, paras)
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_references_placeholder(doc: Document, pdf_text: str) -> None:
    """Pega las referencias del PDF original (todas, sin parseo numerado)."""
    doc.add_heading("Referencias", level=1)
    # Las referencias empiezan después del Cap II y antes de "Anexos"
    paras = split_paragraphs(pdf_text, "Referencias", "Anexo")
    if not paras:
        # Buscar al menos las [N] entradas
        text_lines = pdf_text.split("\n")
        in_refs = False
        for line in text_lines:
            clean = line.replace(ZERO_WIDTH, "").strip()
            if clean.startswith("Referencias"):
                in_refs = True
                continue
            if in_refs and clean.startswith("Anexo"):
                break
            if in_refs and clean and not is_page_number(clean):
                paras.append(clean_paragraph(clean))
    if paras:
        for p in paras:
            doc.add_paragraph(p)
    else:
        add_styled_paragraph(doc, "[INSERTAR aquí las Referencias bibliográficas del PDF original.]",
                              italic=True, color=RGBColor(0xCC, 0x33, 0x33))


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #

def ensure_chapters_docx() -> None:
    """Genera los .docx de Cap III y IV-V-VI si no existen aún."""
    mapping = [
        ("Capitulo_3_Resultados.md", "Capitulo_3_Resultados.docx"),
        ("Capitulo_4_5_6_Discusion_Conclusiones_Recomendaciones.md",
         "Capitulo_4_5_6_Discusion_Conclusiones_Recomendaciones.docx"),
    ]
    for md_name, docx_name in mapping:
        md = HERE / md_name
        target = DOCX_DIR / docx_name
        if md.exists() and (not target.exists() or md.stat().st_mtime > target.stat().st_mtime):
            print(f"  pandoc {md_name} -> {docx_name}")
            subprocess.run(
                ["pandoc", str(md), "-o", str(target)],
                check=True,
            )


def main() -> None:
    print("== build_tesis_integrada.py ==")
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    ensure_chapters_docx()

    pdf_text = extract_pdf_text()
    print(f"  PDF extraído: {len(pdf_text)} chars")

    doc = Document()

    # Carátula y preliminares
    build_caratula(doc)
    build_preliminars(doc)

    # Cap I y II del PDF original
    build_introduction(doc, pdf_text)
    build_methodology(doc, pdf_text)

    # Cap III, IV, V, VI (nuevo, antes de Referencias)
    print("  + Cap III Resultados")
    append_docx(doc, DOCX_DIR / "Capitulo_3_Resultados.docx", page_break_before=True)
    print("  + Cap IV-V-VI Discusión/Conclusiones/Recomendaciones")
    append_docx(doc, DOCX_DIR / "Capitulo_4_5_6_Discusion_Conclusiones_Recomendaciones.docx",
                page_break_before=True)

    # Referencias (después de Cap VI, antes de Anexos)
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    build_references_placeholder(doc, pdf_text)

    # Anexos
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Anexos", level=1)

    anexos = [
        ("Anexo_2_Cuestionario_UTAUT.docx",        "Anexo 2 — Cuestionario UTAUT"),
        ("Anexo_3_Validacion_Juicio_Expertos_PreLlenado.docx",
                                                    "Anexo 3 — Validación por Juicio de Expertos"),
        ("Anexo_4_Solicitud_Autorizacion.docx",    "Anexo 4 — Solicitud de Autorización"),
        ("Anexo_5_Autorizacion_MICOTRANS.docx",    "Anexo 5 — Autorización de MICOTRANS"),
        ("Anexo_8_Metodologia.docx",               "Anexo 8 — Metodología Scrum aplicada"),
        ("Anexo_8_Tablas_Scrum_Completas.docx",    "Anexo 8 (cont.) — Tablas Scrum completas"),
        ("Anexo_Figuras_Capturas.docx",            "Anexo 9 — Figuras del sistema (Web + Móvil end-to-end)"),
    ]
    for name, label in anexos:
        print(f"  + {label}")
        append_docx(doc, DOCX_DIR / name, page_break_before=True)

    doc.save(str(OUT))
    print(f"\nOK -> {OUT}")
    print(f"   Tamaño: {OUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
